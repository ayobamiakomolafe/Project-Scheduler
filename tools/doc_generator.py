from docx import Document

def generate_doc(schedule):
    doc = Document()
    doc.add_heading("Project Schedule Report", 0)

    table = doc.add_table(rows=1, cols=5)
    headers = ["Task", "Duration", "Resources", "Cost", "Dependencies"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for s in schedule:
        row = table.add_row().cells
        row[0].text = s["task"]
        row[1].text = str(s["duration"])
        row[2].text = ", ".join(s["resources"])
        row[3].text = str(s["cost"])
        row[4].text = ", ".join(s["dependencies"])

    doc.save("project.docx")